#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Offline HAZOP participating-expert review assistant."""

from __future__ import annotations

import argparse
import csv
import dataclasses
import datetime as dt
import hashlib
import html
import json
import os
import re
import sys
import zipfile
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET


TEXT_EXTENSIONS = {
    ".txt", ".md", ".csv", ".tsv", ".json", ".yaml", ".yml", ".xml", ".html", ".htm",
    ".log", ".ini", ".cfg", ".dat", ".srt"
}
DOC_EXTENSIONS = {".docx", ".xlsx", ".xlsm", ".xls", ".pdf"}
UNSUPPORTED_EXTENSIONS = {
    ".dwg", ".dxf", ".rvt", ".ifc", ".vsd", ".vsdx", ".jpg", ".jpeg", ".png", ".tif", ".tiff",
    ".bmp", ".gif", ".7z", ".rar", ".zip", ".ppt", ".pptx", ".doc", ".xls"
}
IGNORE_DIR_NAMES = {
    ".git", "__pycache__", "node_modules", ".venv", "venv", ".mypy_cache", ".pytest_cache"
}

PRIORITY_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3}
ISSUE_PREFIX = {"critical": "C", "high": "H", "medium": "M", "low": "L"}
PRIORITY_LABEL = {"critical": "关键", "high": "高", "medium": "中", "low": "低"}
CONFIDENCE_LABEL = {"high": "高", "medium": "中", "low": "低"}

DOC_TYPE_PATTERNS: dict[str, list[str]] = {
    "design_basis": [r"design[_\- ]?basis", r"设计基础", r"设计依据", r"basis of design", r"适用标准", r"标准清单"],
    "pfd": [r"(?<![A-Za-z0-9])PFD(?![A-Za-z0-9])", r"process flow diagram", r"工艺流程图", r"流程图"],
    "pid": [r"(?<![A-Za-z0-9])P\s*&\s*ID(?![A-Za-z0-9])", r"(?<![A-Za-z0-9])P[\-_ ]?ID(?![A-Za-z0-9])", r"piping and instrumentation", r"管道仪表", r"管仪"],
    "equipment_list": [r"equipment[_\- ]?list", r"设备表", r"设备清单", r"equipment datasheet", r"设备数据表"],
    "line_list": [r"line[_\- ]?list", r"管线表", r"管道表", r"line designation", r"管线号"],
    "instrument_index": [r"instrument[_\- ]?index", r"仪表索引", r"仪表清单", r"alarm list", r"报警清单", r"trip list", r"联锁清单"],
    "cause_effect": [r"cause[_\- ]?(?:&|and|\||,)?[_\- ]?effect", r"(?<![A-Za-z0-9])C\s*&\s*E(?![A-Za-z0-9])", r"因果", r"ESD逻辑", r"联锁矩阵"],
    "control_philosophy": [r"control[_\- ]?philosophy", r"控制哲学", r"操作说明", r"operation manual", r"操作手册", r"开车", r"停车"],
    "relief_flare": [r"relief", r"泄压", r"安全阀", r"\bPSV\b", r"\bPRV\b", r"flare", r"火炬", r"vent", r"放空"],
    "hazardous_area": [r"hazardous[_\- ]?area", r"危险区域", r"防爆", r"explosion proof", r"area classification"],
    "layout": [r"plot[_\- ]?plan", r"layout", r"总图", r"设备布置", r"平面布置", r"围堰", r"集液池"],
    "sil_sis": [r"\bSIL\b", r"\bSIS\b", r"安全仪表", r"LOPA", r"保护层"],
    "fire_gas": [r"fire and gas", r"\bFGS\b", r"可燃气", r"火焰探测", r"消防", r"firewater", r"火灾"],
    "material_corrosion": [r"material", r"材料", r"腐蚀", r"corrosion", r"低温材料", r"保冷", r"insulation"],
    "datasheet": [r"datasheet", r"data sheet", r"数据表", r"规格书", r"specification"],
    "incident_history": [r"incident", r"near[_\- ]?miss", r"lesson[_\- ]?learned", r"accident", r"事故", r"未遂", r"险情", r"经验反馈", r"事故教训"],
    "moc_pssr": [r"\bMOC\b", r"management[_\- ]?of[_\- ]?change", r"\bPSSR\b", r"pre[-_ ]?startup", r"变更管理", r"开车前安全审查", r"投产前安全审查"],
    "operating_procedure": [r"operating[_\- ]?procedure", r"\bSOP\b", r"operation[_\- ]?manual", r"操作规程", r"运行规程", r"启动", r"停车", r"应急停车"],
    "maintenance_procedure": [r"maintenance[_\- ]?procedure", r"inspection[_\- ]?test", r"mechanical[_\- ]?integrity", r"检维修", r"维护规程", r"完整性", r"试验记录"],
    "emergency_response": [r"emergency[_\- ]?response", r"emergency[_\- ]?plan", r"fire[_\- ]?drill", r"应急预案", r"消防演练", r"火灾应急"],
    "natural_hazard": [r"natural[_\- ]?hazard", r"flood", r"storm[_\- ]?surge", r"typhoon", r"earthquake", r"lightning", r"自然灾害", r"洪水", r"台风", r"风暴潮", r"地震", r"雷电"],
    "facility_siting": [r"facility[_\- ]?siting", r"stationary[_\- ]?source[_\- ]?siting", r"occupied[_\- ]?building", r"public[_\- ]?receptor", r"building[_\- ]?risk", r"设施选址", r"人员密集", r"公众敏感点", r"建筑物抗爆"],
    "utility_system": [r"utility", r"power[_\- ]?failure", r"standby[_\- ]?power", r"backup[_\- ]?power", r"instrument[_\- ]?air", r"nitrogen", r"cooling[_\- ]?water", r"公用工程", r"备用电源", r"仪表风", r"氮气", r"冷却水"],
    "ship_shore": [r"ship[_\- ]?shore", r"\bSSL\b", r"marine[_\- ]?loading[_\- ]?arm", r"emergency[_\- ]?release", r"\bERS\b", r"船岸", r"装卸臂", r"紧急释放"],
    "cyber_security": [r"cyber", r"network[_\- ]?security", r"\bDCS\b", r"\bPLC\b", r"remote[_\- ]?access", r"网络安全", r"远程访问", r"控制系统安全"],
    "action_tracking": [r"recommendation", r"action[_\- ]?item", r"tracking", r"close[_\- ]?out", r"整改", r"建议项", r"行动项", r"关闭"],
    "hazop_record": [r"HAZOP.*记录", r"HAZOP[_\- ]?worksheet", r"HAZOP[_\- ]?record", r"分析记录表", r"偏差", r"引导词", r"设计意图", r"固有风险", r"残余风险"],
    "high_risk_summary": [r"高风险", r"复杂剧情", r"high[_\- ]?risk", r"complex[_\- ]?scenario"],
    "review_comments": [r"审查意见", r"修改意见", r"review[_\- ]?comment", r"comment[_\- ]?response", r"问题描述", r"修改建议"],
    "sil_report": [r"SIL定级", r"SIL评估", r"安全完整性等级", r"LOPA", r"SIF", r"exSILentia"],
    "sil_record": [r"SIL定级记录", r"保护层分析表", r"Layer of Protection Analysis", r"Hazard Scenario ID", r"RequiredSIL", r"RRF"],
    "interface_boundary": [r"界区", r"界面", r"接口", r"boundary", r"interface", r"battery[_\- ]?limit"],
    "node_diagram": [r"节点图", r"node[_\- ]?diagram"],
    "sds_msds": [r"\bMSDS\b", r"\bSDS\b", r"安全数据表", r"化学品安全技术说明书"],
}

SIGNAL_PATTERNS: dict[str, list[str]] = {
    "lng": [r"\bLNG\b", r"液化天然气", r"接收站", r"卸船", r"BOG", r"boil[- ]?off", r"气化器", r"储罐"],
    "cryogenic": [r"cryogenic", r"低温", r"保冷", r"cold box", r"embrittlement", r"脆裂"],
    "brownfield": [r"扩建", r"改造", r"existing", r"brownfield", r"tie[- ]?in", r"接口", r"切改", r"SIMOPS"],
    "flammable": [r"flammable", r"可燃", r"天然气", r"methane", r"甲烷", r"hydrocarbon", r"烃"],
    "marine_transfer": [r"ship[_\- ]?shore", r"\bSSL\b", r"loading arm", r"marine[_\- ]?loading", r"jetty", r"船岸", r"装卸臂", r"码头"],
    "natural_hazard": [r"natural[_\- ]?hazard", r"flood", r"storm[_\- ]?surge", r"typhoon", r"earthquake", r"lightning", r"自然灾害", r"洪水", r"台风", r"风暴潮", r"地震", r"雷电"],
    "digital_control": [r"\bDCS\b", r"\bSIS\b", r"\bPLC\b", r"cyber", r"remote[_\- ]?access", r"网络安全", r"控制系统", r"远程访问"],
}

TAG_RE = re.compile(
    r"\b(?:"
    r"TK|T|V|P|K|C|E|F|R|H|"
    r"PSV|PRV|TSV|PCV|LCV|FCV|TCV|MOV|XV|SDV|ESDV|BDV|RO|"
    r"PT|PI|PIC|PIT|TT|TI|TIC|TIT|LT|LI|LIC|LIT|FT|FI|FIC|FIT|"
    r"LAH|LAHH|LAL|LALL|PAH|PAHH|PAL|PALL|TAH|TAHH|TAL|TALL|FAL|FAH|"
    r"GD|FD|FG"
    r")[\-_ ]?\d{2,6}[A-Z]?\b",
    re.IGNORECASE,
)


@dataclasses.dataclass
class DocumentRecord:
    path: str
    rel_path: str
    extension: str
    size_bytes: int
    sha256: str
    method: str
    text: str
    doc_types: list[str]
    tags: list[str]
    warnings: list[str]


@dataclasses.dataclass
class Evidence:
    rel_path: str
    keyword: str
    excerpt: str
    tags: list[str]


@dataclasses.dataclass
class Issue:
    issue_id: str
    priority: str
    confidence: str
    topic: str
    node: str
    guidewords: list[str]
    concern: str
    expert_questions: list[str]
    requested_evidence: list[str]
    evidence: list[Evidence]
    source_rule: str


@dataclasses.dataclass
class DrawingFinding:
    finding_id: str
    priority: str
    confidence: str
    parameter: str
    guidewords: list[str]
    analysis_object: str
    problem: str
    possible_causes: list[str]
    possible_consequences: list[str]
    existing_safeguards_to_verify: list[str]
    expert_actions: list[str]
    evidence: list[Evidence]
    related_documents: list[str]
    source_rule: str


@dataclasses.dataclass
class DrawingReview:
    drawing_id: str
    drawing_no: str
    rel_path: str
    title: str
    doc_types: list[str]
    text_chars: int
    extraction_warnings: list[str]
    node_hint: str
    design_intent_hint: str
    review_focus: list[str]
    findings: list[DrawingFinding]


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Review a process design package and draft HAZOP expert questions.")
    parser.add_argument("--input", "-i", required=True, help="Folder containing design files.")
    parser.add_argument("--output", "-o", default=None, help="Folder for review outputs. Defaults to ./hazop_review_output.")
    parser.add_argument("--project-name", default="", help="Project name shown in the report.")
    parser.add_argument("--rule-catalog", default="", help="Optional JSON rule catalog path.")
    parser.add_argument("--max-file-mb", type=float, default=40.0, help="Skip files larger than this size.")
    parser.add_argument("--pdf-page-limit", type=int, default=120, help="Maximum PDF pages to extract per file.")
    parser.add_argument("--max-excerpts-per-rule", type=int, default=4, help="Maximum evidence excerpts per triggered rule.")
    parser.add_argument("--include-unsupported", action="store_true", help="List unsupported binary files in the index.")
    args = parser.parse_args(argv)

    input_dir = Path(args.input).expanduser().resolve()
    if not input_dir.exists() or not input_dir.is_dir():
        print(f"Input folder does not exist or is not a directory: {input_dir}", file=sys.stderr)
        return 2

    output_dir = Path(args.output).expanduser().resolve() if args.output else (Path.cwd() / "hazop_review_output").resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    catalog_path = Path(args.rule_catalog).expanduser().resolve() if args.rule_catalog else default_catalog_path()
    catalog = load_catalog(catalog_path)

    records, unsupported = collect_documents(
        input_dir=input_dir,
        output_dir=output_dir,
        max_file_bytes=int(args.max_file_mb * 1024 * 1024),
        pdf_page_limit=args.pdf_page_limit,
        include_unsupported=args.include_unsupported,
    )
    findings = analyze(records, unsupported, catalog, args.max_excerpts_per_rule)
    package = build_package(
        args.project_name,
        input_dir,
        output_dir,
        records,
        unsupported,
        findings,
        catalog_path,
        catalog,
        max_excerpts_per_rule=args.max_excerpts_per_rule,
    )
    write_outputs(output_dir, package)

    print(f"Wrote HAZOP expert review package to: {output_dir}")
    print(f"Issues: {len(package['issues'])}; documents read: {len(records)}; unsupported/skipped: {len(unsupported)}")
    print(f"Open: {output_dir / 'hazop_expert_opinions.md'}")
    return 0


def default_catalog_path() -> Path:
    return Path(__file__).resolve().parents[1] / "references" / "rule_catalog.json"


def load_catalog(path: Path) -> dict[str, Any]:
    try:
        with path.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as exc:
        raise SystemExit(f"Failed to load rule catalog {path}: {exc}") from exc


def collect_documents(
    input_dir: Path,
    output_dir: Path,
    max_file_bytes: int,
    pdf_page_limit: int,
    include_unsupported: bool,
) -> tuple[list[DocumentRecord], list[dict[str, Any]]]:
    records: list[DocumentRecord] = []
    unsupported: list[dict[str, Any]] = []

    for path in sorted(input_dir.rglob("*")):
        if path.is_dir():
            continue
        if should_ignore(path, input_dir, output_dir):
            continue
        rel = str(path.relative_to(input_dir))
        ext = path.suffix.lower()
        size = path.stat().st_size
        if size > max_file_bytes:
            unsupported.append({"rel_path": rel, "extension": ext, "reason": f"skipped_size_gt_{max_file_bytes}"})
            continue
        if ext not in TEXT_EXTENSIONS and ext not in DOC_EXTENSIONS:
            if include_unsupported or ext in UNSUPPORTED_EXTENSIONS:
                unsupported.append({"rel_path": rel, "extension": ext, "reason": "unsupported_or_binary"})
            continue
        record = extract_document(path, input_dir, pdf_page_limit)
        records.append(record)

    return records, unsupported


def should_ignore(path: Path, input_dir: Path, output_dir: Path) -> bool:
    parts = set(path.relative_to(input_dir).parts[:-1])
    if parts.intersection(IGNORE_DIR_NAMES):
        return True
    try:
        path.resolve().relative_to(output_dir)
        return True
    except ValueError:
        return False


def extract_document(path: Path, input_dir: Path, pdf_page_limit: int) -> DocumentRecord:
    ext = path.suffix.lower()
    warnings: list[str] = []
    method = "text"
    try:
        if ext in TEXT_EXTENSIONS:
            text, method = read_textish(path)
        elif ext == ".docx":
            text, method = read_docx(path)
        elif ext in {".xlsx", ".xlsm"}:
            text, method = read_xlsx(path)
        elif ext == ".xls":
            text, method = read_xls(path)
        elif ext == ".pdf":
            text, method, pdf_warnings = read_pdf(path, pdf_page_limit)
            warnings.extend(pdf_warnings)
        else:
            text, method = "", "unsupported"
    except Exception as exc:
        text = ""
        method = "failed"
        warnings.append(f"extract_failed: {type(exc).__name__}: {exc}")

    if not text.strip():
        warnings.append("no_extractable_text")

    rel = str(path.relative_to(input_dir))
    combined = f"{rel}\n{text[:30000]}"
    doc_types = sorted(classify_doc_types(combined))
    tags = sorted(set(normalize_tag(t) for t in TAG_RE.findall(text)))[:200]
    return DocumentRecord(
        path=str(path),
        rel_path=rel,
        extension=ext,
        size_bytes=path.stat().st_size,
        sha256=sha256_file(path),
        method=method,
        text=normalize_text(text),
        doc_types=doc_types,
        tags=tags,
        warnings=warnings,
    )


def read_textish(path: Path) -> tuple[str, str]:
    raw = path.read_bytes()
    for enc in ("utf-8-sig", "utf-8", "gb18030", "cp936", "cp1252", "latin-1"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode("utf-8", errors="replace")
        enc = "utf-8-replace"
    if path.suffix.lower() in {".html", ".htm"}:
        text = re.sub(r"(?is)<script.*?>.*?</script>", " ", text)
        text = re.sub(r"(?is)<style.*?>.*?</style>", " ", text)
        text = re.sub(r"(?s)<[^>]+>", " ", text)
        text = html.unescape(text)
        return text, f"html:{enc}"
    if path.suffix.lower() == ".csv":
        return csv_to_text(path, enc), f"csv:{enc}"
    if path.suffix.lower() == ".tsv":
        return csv_to_text(path, enc, delimiter="\t"), f"tsv:{enc}"
    return text, f"text:{enc}"


def csv_to_text(path: Path, encoding: str, delimiter: str = ",") -> str:
    rows: list[str] = []
    with path.open("r", encoding=encoding, errors="replace", newline="") as f:
        reader = csv.reader(f, delimiter=delimiter)
        for i, row in enumerate(reader):
            if i >= 5000:
                rows.append("[truncated after 5000 rows]")
                break
            rows.append(" | ".join(str(c) for c in row))
    return "\n".join(rows)


def read_docx(path: Path) -> tuple[str, str]:
    try:
        import docx  # type: ignore

        doc = docx.Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts), "docx:python-docx"
    except Exception:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml")
        root = ET.fromstring(xml)
        texts = [node.text or "" for node in root.iter() if node.tag.endswith("}t")]
        return "\n".join(texts), "docx:xml"


def read_xlsx(path: Path) -> tuple[str, str]:
    try:
        import openpyxl  # type: ignore

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        lines: list[str] = []
        cell_count = 0
        for ws in wb.worksheets:
            lines.append(f"[SHEET] {ws.title}")
            for row in ws.iter_rows(values_only=True):
                values = [str(v) for v in row if v is not None and str(v).strip()]
                if values:
                    lines.append(" | ".join(values))
                cell_count += len(row)
                if cell_count > 120000:
                    lines.append("[truncated after 120000 cells]")
                    return "\n".join(lines), "xlsx:openpyxl:truncated"
        return "\n".join(lines), "xlsx:openpyxl"
    except Exception:
        return read_xlsx_zip_fallback(path), "xlsx:zip-fallback"


def read_xls(path: Path) -> tuple[str, str]:
    try:
        import xlrd  # type: ignore
    except Exception as exc:
        return "", f"xls:xlrd-missing:{type(exc).__name__}"

    book = xlrd.open_workbook(str(path), on_demand=True)
    lines: list[str] = []
    cell_count = 0
    for sheet_name in book.sheet_names():
        sh = book.sheet_by_name(sheet_name)
        lines.append(f"[SHEET] {sheet_name}")
        for ri in range(sh.nrows):
            values = []
            for ci in range(sh.ncols):
                value = sh.cell_value(ri, ci)
                if value not in ("", None):
                    values.append(str(value))
            if values:
                lines.append(" | ".join(values))
            cell_count += sh.ncols
            if cell_count > 120000:
                lines.append("[truncated after 120000 cells]")
                return "\n".join(lines), "xls:xlrd:truncated"
    return "\n".join(lines), "xls:xlrd"


def read_xlsx_zip_fallback(path: Path) -> str:
    texts: list[str] = []
    with zipfile.ZipFile(path) as zf:
        names = [n for n in zf.namelist() if n.endswith(".xml")]
        for name in names:
            if "sharedStrings" not in name and "worksheets" not in name:
                continue
            raw = zf.read(name)
            try:
                root = ET.fromstring(raw)
            except ET.ParseError:
                continue
            for node in root.iter():
                if node.text and node.text.strip():
                    texts.append(node.text.strip())
    return "\n".join(texts[:50000])


def read_pdf(path: Path, page_limit: int) -> tuple[str, str, list[str]]:
    warnings: list[str] = []
    reader = None
    method = ""
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        method = "pdf:pypdf"
    except Exception:
        try:
            from PyPDF2 import PdfReader  # type: ignore

            reader = PdfReader(str(path))
            method = "pdf:PyPDF2"
        except Exception as exc:
            return "", "pdf:failed", [f"pdf_extract_failed: {type(exc).__name__}: {exc}"]

    pages = list(reader.pages)
    if page_limit > 0 and len(pages) > page_limit:
        warnings.append(f"pdf_truncated_pages:{page_limit}_of_{len(pages)}")
        pages = pages[:page_limit]

    parts: list[str] = []
    for i, page in enumerate(pages, start=1):
        try:
            txt = page.extract_text() or ""
        except Exception as exc:
            warnings.append(f"page_{i}_extract_failed:{type(exc).__name__}")
            txt = ""
        if txt.strip():
            parts.append(f"[PAGE {i}]\n{txt}")
    return "\n\n".join(parts), method, warnings


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def classify_doc_types(combined: str) -> set[str]:
    found: set[str] = set()
    for doc_type, patterns in DOC_TYPE_PATTERNS.items():
        if any(re.search(p, combined, flags=re.IGNORECASE) for p in patterns):
            found.add(doc_type)
    return found


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def normalize_tag(tag: str) -> str:
    return re.sub(r"\s+", "-", tag.upper())


def analyze(
    records: list[DocumentRecord],
    unsupported: list[dict[str, Any]],
    catalog: dict[str, Any],
    max_excerpts_per_rule: int,
) -> list[Issue]:
    issues: list[Issue] = []
    doc_counts = doc_type_counts(records)
    signals = detect_signals(records)

    for required in catalog.get("required_document_types", []):
        dtype = required["type"]
        if doc_counts.get(dtype, 0) == 0:
            issues.append(Issue(
                issue_id="",
                priority=required.get("priority", "medium"),
                confidence="high",
                topic=f"资料缺口：{required.get('label', dtype)}",
                node="项目/会议准备",
                guidewords=["Information gap"],
                concern=required.get("why", ""),
                expert_questions=[required.get("question", "请补充该类文件并说明对HAZOP节点的影响。")],
                requested_evidence=[required.get("label", dtype)],
                evidence=[],
                source_rule=f"required_document_type:{dtype}",
            ))

    if not records:
        issues.append(Issue(
            issue_id="",
            priority="critical",
            confidence="high",
            topic="无法读取设计资料",
            node="项目/会议准备",
            guidewords=["Information gap"],
            concern="输入目录没有可读取的文本、Word、Excel或PDF内容，无法形成有证据的专家问题。",
            expert_questions=["请提供可复制文本的PDF、Word、Excel、P&ID导出文本、OCR结果或图纸索引；哪些关键图纸目前只有扫描件？"],
            requested_evidence=["可读取的设计资料包", "图纸目录", "OCR文本"],
            evidence=[],
            source_rule="system:no_records",
        ))

    unreadable = [r for r in records if "no_extractable_text" in r.warnings]
    if unreadable:
        sample = [Evidence(r.rel_path, "no_extractable_text", "文件可见但未能提取文本，可能是扫描件或二进制图纸。", []) for r in unreadable[:8]]
        issues.append(Issue(
            issue_id="",
            priority="medium",
            confidence="high",
            topic="部分文件不可文本检索",
            node="资料质量",
            guidewords=["Information gap"],
            concern="扫描件、图片型PDF或原生CAD文件未OCR时，工具无法检查其中的阀门、仪表、节点和注释。",
            expert_questions=["请确认这些文件是否为最新有效设计文件；如是，请提供OCR文本、可搜索PDF或图纸数据导出。"],
            requested_evidence=["OCR结果", "可搜索PDF", "图纸清单"],
            evidence=sample,
            source_rule="system:unreadable_files",
        ))

    if unsupported:
        sample = [Evidence(u["rel_path"], u.get("extension", ""), f"未解析：{u.get('reason', '')}", []) for u in unsupported[:10]]
        issues.append(Issue(
            issue_id="",
            priority="low",
            confidence="high",
            topic="存在未解析的二进制或大文件",
            node="资料质量",
            guidewords=["Information gap"],
            concern="CAD、图片、压缩包或超大文件可能包含HAZOP关键信息，但本次离线文本引擎未解析。",
            expert_questions=["这些未解析文件中是否包含P&ID、总图、危险区域、防火、供应商数据或C&E矩阵？请提供对应可搜索版本。"],
            requested_evidence=["文件目录说明", "可搜索导出版"],
            evidence=sample,
            source_rule="system:unsupported_files",
        ))

    for rule in catalog.get("rules", []):
        if not rule_applies(rule, records, signals):
            continue
        evidence = find_evidence(records, rule.get("any_keywords", []), max_excerpts_per_rule)
        confidence = "medium" if evidence else "low"
        if evidence and len(evidence) >= min(2, max_excerpts_per_rule):
            confidence = "high"
        issues.append(Issue(
            issue_id="",
            priority=rule.get("priority", "medium"),
            confidence=confidence,
            topic=rule.get("topic", rule.get("id", "HAZOP问题")),
            node=infer_node(rule, evidence),
            guidewords=rule.get("guidewords", []),
            concern=rule.get("concern", ""),
            expert_questions=rule.get("questions", []),
            requested_evidence=rule.get("requested_evidence", []),
            evidence=evidence,
            source_rule=rule.get("id", "rule"),
        ))

    issues = sorted(issues, key=lambda x: (PRIORITY_ORDER.get(x.priority, 9), x.topic, x.source_rule))
    for index, issue in enumerate(issues, start=1):
        issue.issue_id = f"HAZOP-{ISSUE_PREFIX.get(issue.priority, 'X')}{index:03d}"
    return issues


def build_drawing_reviews(
    records: list[DocumentRecord],
    catalog: dict[str, Any],
    signals: set[str],
    max_findings_per_drawing: int = 8,
    max_evidence_per_finding: int = 3,
) -> list[DrawingReview]:
    reviews: list[DrawingReview] = []
    reviewable = [r for r in records if is_reviewable_drawing(r)]
    for drawing_index, record in enumerate(reviewable, start=1):
        drawing_id = f"DWG-{drawing_index:03d}"
        findings: list[DrawingFinding] = []
        for rule in catalog.get("rules", []):
            if not rule_applies_to_record(rule, record, signals):
                continue
            evidence = find_evidence_in_record(record, rule.get("any_keywords", []), max_evidence_per_finding)
            if not evidence:
                continue
            finding = make_drawing_finding(
                drawing_id=drawing_id,
                serial=len(findings) + 1,
                record=record,
                rule=rule,
                evidence=evidence,
                related_documents=find_related_documents(record, records, rule, evidence),
            )
            findings.append(finding)

        findings = sorted(
            findings,
            key=lambda item: (PRIORITY_ORDER.get(item.priority, 9), item.parameter, item.source_rule),
        )[:max_findings_per_drawing]
        for serial, finding in enumerate(findings, start=1):
            finding.finding_id = f"{drawing_id}-F{serial:02d}"

        reviews.append(DrawingReview(
            drawing_id=drawing_id,
            drawing_no=infer_drawing_no(record),
            rel_path=record.rel_path,
            title=infer_drawing_title(record),
            doc_types=record.doc_types,
            text_chars=len(record.text),
            extraction_warnings=record.warnings,
            node_hint=infer_drawing_node(record),
            design_intent_hint=infer_design_intent(record),
            review_focus=infer_review_focus(record),
            findings=findings,
        ))
    return reviews


def is_reviewable_drawing(record: DocumentRecord) -> bool:
    if is_existing_drawing(record.rel_path) or is_standard_reference_drawing(record.rel_path):
        return False
    if is_pfd_context_drawing(record):
        return False
    if "no_extractable_text" in record.warnings:
        return is_pid_drawing(record)
    return is_pid_drawing(record)


def is_existing_drawing(rel_path: str) -> bool:
    return "已建" in rel_path


def is_standard_reference_drawing(rel_path: str) -> bool:
    tokens = ("通用注释", "通用图例", "通用管道详图", "通用仪表详图", "通用仪表缩写", "取样详图")
    return any(token in rel_path for token in tokens)


def is_pfd_context_drawing(record: DocumentRecord) -> bool:
    name = record.rel_path.lower()
    if "pid" in record.doc_types or "p&id" in name or "管道及仪表" in record.rel_path or "管道仪表" in record.rel_path:
        return False
    return "pfd" in record.doc_types or "pfd" in name or "工艺流程图" in record.rel_path


def is_pid_drawing(record: DocumentRecord) -> bool:
    name = record.rel_path.lower()
    pid_tokens = ("p&id", "pid", "管道及仪表流程图", "管道仪表流程图")
    return any(token in name or token in record.rel_path for token in pid_tokens)


def rule_applies_to_record(rule: dict[str, Any], record: DocumentRecord, signals: set[str]) -> bool:
    required_signals = set(rule.get("signals", []))
    if required_signals and not required_signals.issubset(signals):
        return False
    doc_types = set(rule.get("doc_types", []))
    if doc_types and not doc_types.intersection(record.doc_types):
        return False
    keywords = rule.get("any_keywords", [])
    if not keywords:
        return False
    search_space = f"{record.rel_path}\n{record.text[:200000]}"
    return any(re.search(re.escape(k), search_space, flags=re.IGNORECASE) for k in keywords)


def find_evidence_in_record(record: DocumentRecord, keywords: Iterable[str], limit: int) -> list[Evidence]:
    evidence: list[Evidence] = []
    seen: set[str] = set()
    search_space = f"{record.rel_path}\n{record.text}"
    for keyword in keywords:
        if len(evidence) >= limit:
            break
        if keyword.lower() in seen:
            continue
        pattern = re.compile(re.escape(keyword), flags=re.IGNORECASE)
        match = pattern.search(search_space)
        if not match:
            continue
        seen.add(keyword.lower())
        excerpt = make_excerpt(search_space, match.start(), match.end())
        evidence.append(Evidence(record.rel_path, keyword, excerpt, sorted(set(normalize_tag(t) for t in TAG_RE.findall(excerpt)))[:10]))
    return evidence


def make_drawing_finding(
    drawing_id: str,
    serial: int,
    record: DocumentRecord,
    rule: dict[str, Any],
    evidence: list[Evidence],
    related_documents: list[str],
) -> DrawingFinding:
    questions = rule.get("questions", [])
    requested_evidence = rule.get("requested_evidence", [])
    topic = rule.get("topic", rule.get("id", "HAZOP问题"))
    guidewords = rule.get("guidewords", [])
    confidence = "high" if len(evidence) >= 2 else "medium"
    safeguards = infer_safeguards_to_verify(record, evidence, rule)
    actions = []
    if questions:
        actions.extend(questions[:2])
    if requested_evidence:
        actions.append("会上要求定位或补充：" + "、".join(str(x) for x in requested_evidence[:5]))
    if related_documents:
        actions.append("必要时跨图核查：" + "；".join(related_documents[:4]))

    return DrawingFinding(
        finding_id=f"{drawing_id}-F{serial:02d}",
        priority=rule.get("priority", "medium"),
        confidence=confidence,
        parameter=infer_parameter(rule),
        guidewords=guidewords,
        analysis_object=infer_analysis_object(record, evidence, topic),
        problem=questions[0] if questions else f"请核查本图纸中与“{topic}”相关的偏离是否已进入HAZOP讨论。",
        possible_causes=infer_possible_causes(rule, guidewords),
        possible_consequences=[rule.get("concern", "") or "该偏离的后果需要结合本图纸、上下游接口和保护层进一步确认。"],
        existing_safeguards_to_verify=safeguards,
        expert_actions=actions or ["请设计方在会上定位图纸、节点、回路和证据，并说明是否需要形成HAZOP行动项。"],
        evidence=evidence,
        related_documents=related_documents,
        source_rule=rule.get("id", "rule"),
    )


def infer_parameter(rule: dict[str, Any]) -> str:
    guidewords = " ".join(str(x).lower() for x in rule.get("guidewords", []))
    topic = str(rule.get("topic", "")).lower()
    if any(x in guidewords or x in topic for x in ("membrane", "containment", "insulation", "薄膜", "次屏障", "绝热", "保冷")):
        return "围护/泄漏监测"
    if any(x in guidewords or x in topic for x in ("rollover", "stratification", "composition", "翻滚", "分层", "密度", "ltd")):
        return "组分/分层"
    if any(x in guidewords or x in topic for x in ("bog", "recondenser", "boil", "再冷凝")):
        return "BOG/压力"
    if any(x in guidewords or x in topic for x in ("ship", "shore", "jetty", "vapor return", "船岸", "卸船", "装船", "返气")):
        return "船岸/转输"
    if any(x in guidewords or x in topic for x in ("pressure", "vacuum", "压力", "泄压", "超压")):
        return "压力"
    if any(x in guidewords or x in topic for x in ("temperature", "low temperature", "温度", "低温", "预冷")):
        return "温度"
    if any(x in guidewords or x in topic for x in ("flow", "reverse", "流量", "反流", "液击")):
        return "流量/流向"
    if any(x in guidewords or x in topic for x in ("level", "液位", "储罐")):
        return "液位/储量"
    if any(x in guidewords or x in topic for x in ("control", "alarm", "interlock", "sis", "esd", "联锁", "报警")):
        return "控制/报警/联锁"
    if any(x in guidewords or x in topic for x in ("utility", "power", "公用工程", "仪表风", "氮气")):
        return "公用工程"
    if any(x in guidewords or x in topic for x in ("maintenance", "startup", "shutdown", "检修", "开车", "停车")):
        return "操作/检维修状态"
    return "设计意图/接口"


def infer_analysis_object(record: DocumentRecord, evidence: list[Evidence], topic: str) -> str:
    tags: list[str] = []
    for ev in evidence:
        tags.extend(ev.tags)
    if tags:
        return "、".join(sorted(set(tags))[:6])
    if record.tags:
        return "、".join(record.tags[:6])
    return topic


def infer_possible_causes(rule: dict[str, Any], guidewords: list[str]) -> list[str]:
    text = " ".join([str(rule.get("topic", "")), str(rule.get("id", "")), " ".join(guidewords)]).lower()
    causes: list[str] = []
    if any(x in text for x in ("membrane", "containment", "insulation", "薄膜", "次屏障", "绝热", "保冷")):
        causes.append("主膜局部缺陷、穿罐管口热位移、绝热空间氮封/差压异常、气体检测失效、保冷受潮或低温冷点未被及时识别。")
    if any(x in text for x in ("rollover", "stratification", "composition", "翻滚", "分层", "密度", "ltd")):
        causes.append("多气源密度差、顶部/底部进液策略不当、长期低周转储存、倒罐/返输扰动、LTD趋势报警或混合循环策略不足。")
    if any(x in text for x in ("bog", "recondenser", "boil", "再冷凝")):
        causes.append("卸船或翻滚导致BOG突增、压缩机跳车、再冷凝器液位/压力控制失效、高压泵或外输负荷突变、火炬背压限制。")
    if any(x in text for x in ("ship", "shore", "jetty", "vapor return", "船岸", "卸船", "装船", "返气")):
        causes.append("船岸ESD时序不一致、返气不足或误隔离、装卸臂位移/ERS动作、保冷未完成、装船/卸船模式阀位冲突。")
    if any(x in text for x in ("orv", "seawater", "海水", "气化器", "冷排水")):
        causes.append("海水泵跳车、滤网或海生物堵塞、加氯异常、ORV结冰、出口低温联锁或外输负荷切换处理不足。")
    if any(x in text for x in ("flare", "relief", "thermal expansion", "火炬", "放空", "热膨胀", "tsv")):
        causes.append("封闭低温液段受热、TSV/PSV覆盖边界不清、低温两相夹液进入放空、分液罐或火炬总管最低温度/背压不满足。")
    if any(x in text for x in ("no flow", "less flow", "reverse", "flow", "反流", "流量")):
        causes.append("阀门误关/误开、泵停、止回阀失效、旁路打开、过滤或管线堵塞、上下游压差变化。")
    if any(x in text for x in ("pressure", "blocked", "vacuum", "泄压", "超压", "封闭")):
        causes.append("封闭液段受热、下游切断、控制阀失效、泄放路径受限、火灾或外部热输入。")
    if any(x in text for x in ("temperature", "low temperature", "cryogenic", "低温", "预冷")):
        causes.append("低温介质泄漏、预冷/吹扫介质误入、保冷失效、冷量传递至非低温材料或结构。")
    if any(x in text for x in ("control", "alarm", "interlock", "sis", "esd", "联锁", "报警")):
        causes.append("仪表故障、设定值待定、旁路/屏蔽、复位条件不清、公用工程或通信失效。")
    if any(x in text for x in ("maintenance", "startup", "shutdown", "moc", "pssr", "检修", "开车", "停车", "切改")):
        causes.append("临时盲板/短接/旁路、置换不充分、隔离边界错误、投产切换顺序不清。")
    if not causes:
        causes.append("本图纸出现相关对象或接口，具体原因需由设计方结合图纸、操作规程和上下游资料确认。")
    return causes


def infer_safeguards_to_verify(record: DocumentRecord, evidence: list[Evidence], rule: dict[str, Any]) -> list[str]:
    search_text = " ".join([record.text[:30000], " ".join(ev.excerpt for ev in evidence)])
    candidates = [
        ("ESD", "ESD/紧急停车动作"),
        ("SIS", "SIS/安全仪表功能"),
        ("GDS", "GDS/可燃气检测系统"),
        ("FGS", "FGS/火气系统"),
        ("FAS", "FAS/火灾报警系统"),
        ("PSV", "PSV/压力安全阀"),
        ("PRV", "PRV/泄压阀"),
        ("TSV", "TSV/热膨胀安全阀"),
        ("XV", "XV/切断阀"),
        ("LTD", "LTD/液位温度密度监测"),
        ("氧含量", "氧含量监测"),
        ("可燃气", "可燃气检测"),
        ("温度", "温度/冷点监测"),
        ("压力", "压力报警或控制"),
        ("液位", "液位报警或联锁"),
        ("BOG", "BOG压力控制/压缩处理"),
        ("再冷凝", "BOG再冷凝能力"),
        ("最小流量", "泵最小流量保护"),
        ("回流", "回流/防反流路径"),
        ("报警", "报警与操作员响应"),
        ("联锁", "联锁逻辑"),
        ("止回阀", "止回阀/防反流措施"),
        ("氮气", "氮气吹扫/置换"),
        ("氮气缓冲", "氮气缓冲和调压"),
        ("放空", "放空/火炬去向"),
    ]
    found = [label for token, label in candidates if re.search(re.escape(token), search_text, flags=re.IGNORECASE)]
    if found:
        return [f"图纸可抽取文本出现“{item}”，需核查其是否为本场景有效保护层及是否进入C&E/操作规程。" for item in found[:5]]
    if "no_extractable_text" in record.warnings:
        return ["该图纸未能抽取文本，需人工看图确认已有保护层、阀位、联锁和注释。"]
    return ["可抽取文本未明确识别已有保护层，需人工看图并结合C&E、阀门表、仪表索引确认。"]


def find_related_documents(
    record: DocumentRecord,
    records: list[DocumentRecord],
    rule: dict[str, Any],
    evidence: list[Evidence],
    limit: int = 5,
) -> list[str]:
    topic = str(rule.get("topic", ""))
    rule_id = str(rule.get("id", ""))
    if not re.search(r"接口|ESD|SIS|C&E|BOG|公用工程|船岸|TRANSFER|UTILITY|INTERFACE|TRACE", topic + rule_id, re.IGNORECASE):
        return []
    tags: set[str] = set(record.tags[:20])
    for ev in evidence:
        tags.update(ev.tags)
    keywords = [str(k) for k in rule.get("any_keywords", [])[:4]]
    related: list[str] = []
    for other in records:
        if other.rel_path == record.rel_path:
            continue
        score = 0
        if tags and tags.intersection(other.tags):
            score += 2
        other_space = f"{other.rel_path}\n{other.text[:50000]}"
        if any(re.search(re.escape(k), other_space, flags=re.IGNORECASE) for k in keywords):
            score += 1
        if score > 0:
            related.append(other.rel_path)
        if len(related) >= limit:
            break
    return related


def infer_drawing_no(record: DocumentRecord) -> str:
    stem = Path(record.rel_path).stem
    match = re.match(r"([A-Za-z0-9#_.\-]+)", stem)
    return match.group(1) if match else stem


def infer_drawing_title(record: DocumentRecord) -> str:
    stem = Path(record.rel_path).stem
    parts = re.split(r"\s+", stem, maxsplit=1)
    return parts[1] if len(parts) > 1 else stem


def infer_drawing_node(record: DocumentRecord) -> str:
    name = record.rel_path
    mapping = [
        ("薄膜罐", "薄膜LNG储罐氮气/绝热空间"),
        ("储罐", "LNG储罐系统"),
        ("低压输送泵", "LNG低压输送泵系统"),
        ("低压泵", "LNG低压输送泵系统"),
        ("装船泵", "LNG装船泵系统"),
        ("BOG", "BOG压缩/回收系统"),
        ("火炬", "火炬/放空系统"),
        ("ORV", "ORV气化/外输接口"),
        ("仪表空气", "仪表空气系统"),
        ("工厂空气", "工厂空气系统"),
        ("氮气", "氮气分配/吹扫系统"),
        ("卸船", "LNG卸船/码头接口"),
        ("槽车", "低压槽车外输系统"),
    ]
    for token, node in mapping:
        if token.lower() in name.lower():
            return node
    if record.tags:
        return "相关设备/仪表：" + "、".join(record.tags[:6])
    if "pid" in record.doc_types:
        return "P&ID图纸节点待专家划分"
    if "pfd" in record.doc_types:
        return "PFD系统边界待专家划分"
    return "图纸/资料节点待专家确认"


def infer_design_intent(record: DocumentRecord) -> str:
    node = infer_drawing_node(record)
    if "PFD" in node or "pfd" in record.doc_types:
        return "确认工艺路径、主要设备、正常工况、物料去向和系统边界，作为HAZOP节点设计意图的上游依据。"
    if "薄膜" in node:
        return "维持薄膜罐绝热/围护空间氮气惰化、压力/差压、泄漏监测、排放和异常处置在厂家与项目设计包络内。"
    if "储罐" in node:
        return "维持LNG储罐液位、压力、真空保护、BOG排出、分层翻滚监测和预冷/放空/吹扫路径在设计包络内。"
    if "低压输送泵" in node:
        return "从LNG储罐按要求输送LNG，避免无流、反流、液击、低温泄漏和泵保护失效。"
    if "装船泵" in node:
        return "按装船工况输送LNG，并与返气、ESD、预冷、放空和接口切断保持一致。"
    if "BOG" in node:
        return "收集、压缩或处置BOG，维持储罐压力和下游系统压力在允许范围内。"
    if "火炬" in node or "放空" in node:
        return "接收安全阀、放空和事故排放，保证排放去向、背压和二次风险可控。"
    if "空气" in node or "氮气" in node:
        return "向相关用户稳定供应公用工程，并明确失效时阀位、联锁和安全状态。"
    if "pid" in record.doc_types:
        return "逐项核查本图纸中的管线、阀门、仪表、联锁、泄压、放空、排凝、吹扫和跨图接口。"
    return "确认该资料对HAZOP节点、设计意图、偏离原因、后果和保护层的支持程度。"


def infer_review_focus(record: DocumentRecord) -> list[str]:
    focus = ["节点边界和设计意图", "偏离、原因、后果、已有措施和需补证据"]
    text = record.rel_path + "\n" + record.text[:20000]
    candidates = [
        (r"LNG|低温|储罐|BOG", "LNG低温、储罐压力/液位和BOG去向"),
        (r"ESD|SIS|联锁|报警|XV|SDV", "ESD/SIS/报警联锁和阀门故障位"),
        (r"PSV|PRV|TSV|安全阀|泄压|放空|火炬", "泄压、放空、火炬和封闭液段热膨胀"),
        (r"氮气|仪表空气|工厂空气|公用工程", "公用工程失效和共同原因失效"),
        (r"接口|界区|已建|改造|扩建|预留", "新老系统接口、切改和跨图关系"),
        (r"预冷|开车|停车|检修|盲板|短接|置换", "预冷、开停车、检维修和临时状态"),
    ]
    for pattern, item in candidates:
        if re.search(pattern, text, flags=re.IGNORECASE):
            focus.append(item)
    return list(dict.fromkeys(focus))[:7]


def doc_type_counts(records: list[DocumentRecord]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for r in records:
        for dtype in r.doc_types:
            counts[dtype] = counts.get(dtype, 0) + 1
    return counts


def detect_signals(records: list[DocumentRecord]) -> set[str]:
    text = "\n".join(f"{r.rel_path}\n{r.text[:100000]}" for r in records)
    signals: set[str] = set()
    for signal, patterns in SIGNAL_PATTERNS.items():
        if any(re.search(p, text, flags=re.IGNORECASE) for p in patterns):
            signals.add(signal)
    return signals


def rule_applies(rule: dict[str, Any], records: list[DocumentRecord], signals: set[str]) -> bool:
    required_signals = set(rule.get("signals", []))
    if required_signals and not required_signals.issubset(signals):
        return False
    doc_types = set(rule.get("doc_types", []))
    if doc_types:
        found_types = set()
        for r in records:
            found_types.update(r.doc_types)
        if not doc_types.intersection(found_types):
            return False
    keywords = rule.get("any_keywords", [])
    if not keywords:
        return True
    haystack = "\n".join(f"{r.rel_path}\n{r.text[:200000]}" for r in records)
    return any(re.search(re.escape(k), haystack, flags=re.IGNORECASE) for k in keywords)


def find_evidence(records: list[DocumentRecord], keywords: Iterable[str], limit: int) -> list[Evidence]:
    evidence: list[Evidence] = []
    seen: set[tuple[str, str]] = set()
    for keyword in keywords:
        if len(evidence) >= limit:
            break
        pattern = re.compile(re.escape(keyword), flags=re.IGNORECASE)
        for record in records:
            if len(evidence) >= limit:
                break
            search_space = f"{record.rel_path}\n{record.text}"
            match = pattern.search(search_space)
            if not match:
                continue
            key = (record.rel_path, keyword.lower())
            if key in seen:
                continue
            seen.add(key)
            excerpt = make_excerpt(search_space, match.start(), match.end())
            evidence.append(Evidence(record.rel_path, keyword, excerpt, sorted(set(normalize_tag(t) for t in TAG_RE.findall(excerpt)))[:10]))
    return evidence


def make_excerpt(text: str, start: int, end: int, width: int = 460) -> str:
    left = max(0, start - width // 2)
    right = min(len(text), end + width // 2)
    excerpt = text[left:right]
    excerpt = re.sub(r"\s+", " ", excerpt).strip()
    if left > 0:
        excerpt = "..." + excerpt
    if right < len(text):
        excerpt += "..."
    return excerpt


def infer_node(rule: dict[str, Any], evidence: list[Evidence]) -> str:
    topic = rule.get("topic", "")
    if evidence:
        tags = []
        for ev in evidence:
            tags.extend(ev.tags)
        if tags:
            return "相关设备/仪表：" + ", ".join(sorted(set(tags))[:8])
    if "LNG" in topic or "低温" in topic:
        return "LNG/低温系统"
    if "BOG" in topic:
        return "BOG系统"
    if "储罐" in topic:
        return "储罐系统"
    if "气化" in topic:
        return "气化/外输系统"
    return "项目/系统层面"


def build_package(
    project_name: str,
    input_dir: Path,
    output_dir: Path,
    records: list[DocumentRecord],
    unsupported: list[dict[str, Any]],
    issues: list[Issue],
    catalog_path: Path,
    catalog: dict[str, Any],
    max_excerpts_per_rule: int,
) -> dict[str, Any]:
    counts = doc_type_counts(records)
    signal_set = detect_signals(records)
    signals = sorted(signal_set)
    project_context = build_project_context(records)
    drawing_reviews = build_drawing_reviews(
        records=records,
        catalog=catalog,
        signals=signal_set,
        max_findings_per_drawing=8,
        max_evidence_per_finding=max(1, min(max_excerpts_per_rule, 3)),
    )
    now = dt.datetime.now().astimezone().isoformat(timespec="seconds")
    return {
        "metadata": {
            "project_name": project_name or input_dir.name,
            "generated_at": now,
            "input_dir": str(input_dir),
            "output_dir": str(output_dir),
            "rule_catalog": str(catalog_path),
            "role": "HAZOP参会专家，不是主席",
            "method_note": "离线证据筛查；问题必须经人类专家看图复核后才能用于会议。",
        },
        "signals": signals,
        "project_context": project_context,
        "practice_sources": catalog.get("practice_sources", []),
        "document_type_counts": counts,
        "documents": [record_to_index(r) for r in records],
        "unsupported_or_skipped": unsupported,
        "issues": [issue_to_dict(i) for i in issues],
        "drawing_reviews": [drawing_review_to_dict(r) for r in drawing_reviews],
    }


def build_project_context(records: list[DocumentRecord]) -> dict[str, Any]:
    pfd_records = [r for r in records if is_pfd_context_drawing(r)]
    ignored_existing = [r for r in records if is_existing_drawing(r.rel_path)]
    ignored_reference = [
        r for r in records
        if is_standard_reference_drawing(r.rel_path) and not is_existing_drawing(r.rel_path)
    ]
    return {
        "pfd_context_documents": [
            {
                "rel_path": r.rel_path,
                "drawing_no": infer_drawing_no(r),
                "title": infer_drawing_title(r),
                "text_chars": len(r.text),
                "tags_sample": r.tags[:20],
            }
            for r in pfd_records
        ],
        "ignored_existing_drawings": [
            {
                "rel_path": r.rel_path,
                "drawing_no": infer_drawing_no(r),
                "title": infer_drawing_title(r),
                "reason": "文件名或路径含“已建”，默认作为既有装置接口背景，不进入本轮逐张P&ID预审。",
            }
            for r in ignored_existing
        ],
        "ignored_reference_drawings": [
            {
                "rel_path": r.rel_path,
                "drawing_no": infer_drawing_no(r),
                "title": infer_drawing_title(r),
                "reason": "通用注释、图例或详图不是单张工艺节点P&ID，默认不生成逐图纸问题行。",
            }
            for r in ignored_reference
        ],
        "review_scope_note": "逐图纸分析默认只覆盖新建/扩建P&ID；PFD用于理解项目整体工艺路径和系统边界；已建图纸和通用参考图默认忽略。",
    }


def record_to_index(record: DocumentRecord) -> dict[str, Any]:
    return {
        "rel_path": record.rel_path,
        "extension": record.extension,
        "size_bytes": record.size_bytes,
        "sha256": record.sha256,
        "method": record.method,
        "doc_types": record.doc_types,
        "tag_count": len(record.tags),
        "tags_sample": record.tags[:40],
        "text_chars": len(record.text),
        "warnings": record.warnings,
    }


def issue_to_dict(issue: Issue) -> dict[str, Any]:
    return {
        "issue_id": issue.issue_id,
        "priority": issue.priority,
        "confidence": issue.confidence,
        "topic": issue.topic,
        "node": issue.node,
        "guidewords": issue.guidewords,
        "concern": issue.concern,
        "expert_questions": issue.expert_questions,
        "requested_evidence": issue.requested_evidence,
        "evidence": [dataclasses.asdict(e) for e in issue.evidence],
        "source_rule": issue.source_rule,
    }


def drawing_review_to_dict(review: DrawingReview) -> dict[str, Any]:
    return {
        "drawing_id": review.drawing_id,
        "drawing_no": review.drawing_no,
        "rel_path": review.rel_path,
        "title": review.title,
        "doc_types": review.doc_types,
        "text_chars": review.text_chars,
        "extraction_warnings": review.extraction_warnings,
        "node_hint": review.node_hint,
        "design_intent_hint": review.design_intent_hint,
        "review_focus": review.review_focus,
        "findings": [drawing_finding_to_dict(f) for f in review.findings],
    }


def drawing_finding_to_dict(finding: DrawingFinding) -> dict[str, Any]:
    return {
        "finding_id": finding.finding_id,
        "priority": finding.priority,
        "confidence": finding.confidence,
        "parameter": finding.parameter,
        "guidewords": finding.guidewords,
        "analysis_object": finding.analysis_object,
        "problem": finding.problem,
        "possible_causes": finding.possible_causes,
        "possible_consequences": finding.possible_consequences,
        "existing_safeguards_to_verify": finding.existing_safeguards_to_verify,
        "expert_actions": finding.expert_actions,
        "evidence": [dataclasses.asdict(e) for e in finding.evidence],
        "related_documents": finding.related_documents,
        "source_rule": finding.source_rule,
    }


def write_outputs(output_dir: Path, package: dict[str, Any]) -> None:
    (output_dir / "hazop_expert_opinions.json").write_text(
        json.dumps(package, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (output_dir / "extracted_index.json").write_text(
        json.dumps({
            "metadata": package["metadata"],
            "signals": package["signals"],
            "project_context": package.get("project_context", {}),
            "practice_sources": package.get("practice_sources", []),
            "document_type_counts": package["document_type_counts"],
            "documents": package["documents"],
            "unsupported_or_skipped": package["unsupported_or_skipped"],
        }, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (output_dir / "hazop_expert_opinions.md").write_text(render_markdown(package), encoding="utf-8")
    write_csv(output_dir / "hazop_expert_questions.csv", package["issues"])
    (output_dir / "hazop_drawing_reviews.json").write_text(
        json.dumps({
            "metadata": package["metadata"],
            "signals": package["signals"],
            "project_context": package.get("project_context", {}),
            "drawing_reviews": package["drawing_reviews"],
        }, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (output_dir / "hazop_drawing_reviews.md").write_text(render_drawing_reviews_markdown(package), encoding="utf-8")
    write_drawing_review_csv(output_dir / "hazop_drawing_reviews.csv", package["drawing_reviews"])
    (output_dir / "review_prompt.md").write_text(render_review_prompt(package), encoding="utf-8")
    (output_dir / "run_summary.json").write_text(
        json.dumps({
            "generated_at": package["metadata"]["generated_at"],
            "project_name": package["metadata"]["project_name"],
            "issue_count": len(package["issues"]),
            "drawing_review_count": len(package["drawing_reviews"]),
            "drawing_finding_count": sum(len(r["findings"]) for r in package["drawing_reviews"]),
            "pfd_context_count": len(package.get("project_context", {}).get("pfd_context_documents", [])),
            "ignored_existing_count": len(package.get("project_context", {}).get("ignored_existing_drawings", [])),
            "ignored_reference_count": len(package.get("project_context", {}).get("ignored_reference_drawings", [])),
            "document_count": len(package["documents"]),
            "unsupported_or_skipped_count": len(package["unsupported_or_skipped"]),
            "outputs": [
                "hazop_drawing_reviews.md",
                "hazop_drawing_reviews.json",
                "hazop_drawing_reviews.csv",
                "hazop_expert_opinions.md",
                "hazop_expert_opinions.json",
                "hazop_expert_questions.csv",
                "extracted_index.json",
                "review_prompt.md",
            ],
        }, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def render_markdown(package: dict[str, Any]) -> str:
    meta = package["metadata"]
    issues = package["issues"]
    context = package.get("project_context", {})
    priority_counts: dict[str, int] = {}
    for issue in issues:
        priority_counts[issue["priority"]] = priority_counts.get(issue["priority"], 0) + 1

    lines: list[str] = []
    lines.append(f"# HAZOP专家意见筛查报告 - {meta['project_name']}")
    lines.append("")
    lines.append(f"- 生成时间：{meta['generated_at']}")
    lines.append(f"- 输入目录：`{meta['input_dir']}`")
    lines.append(f"- 工具角色：HAZOP参会专家问题筛查，不是主席结论或正式会议纪要")
    lines.append(f"- 识别项数：{len(issues)}")
    lines.append(f"- 优先级统计：{', '.join(f'{PRIORITY_LABEL.get(k, k)}={v}' for k, v in sorted(priority_counts.items(), key=lambda x: PRIORITY_ORDER.get(x[0], 9))) or '无'}")
    lines.append("")
    lines.append("## 使用边界")
    lines.append("")
    lines.append("本报告用于会前准备和专家发问。它依据输入文件的可提取文本和规则库生成问题，不代表设计确认、风险评级、主席行动项关闭或正式HAZOP worksheet。扫描图、CAD原图、压缩包和无法提取文本的PDF需要另行OCR或人工核查。")
    lines.append("")
    lines.append("本轮逐张图纸分析默认只覆盖新建/扩建P&ID。PFD只作为理解项目整体工艺路径、系统边界和上下游关系的基础，不生成逐图纸问题行；文件名或路径含“已建”的图纸默认作为既有装置接口背景，不进入本轮逐张分析。")
    if context:
        lines.append("")
        lines.append("## 项目理解基础与默认忽略范围")
        lines.append("")
        lines.append(f"- PFD背景图纸：{len(context.get('pfd_context_documents', []))} 份")
        lines.append(f"- 默认忽略已建图纸：{len(context.get('ignored_existing_drawings', []))} 份")
        lines.append(f"- 默认忽略通用参考图：{len(context.get('ignored_reference_drawings', []))} 份")
        for item in context.get("pfd_context_documents", [])[:6]:
            lines.append(f"- PFD背景：`{item.get('rel_path', '')}`")
    lines.append("")
    if package.get("practice_sources"):
        lines.append("## 借鉴来源")
        lines.append("")
        for source in package["practice_sources"]:
            label = source.get("label", "")
            url = source.get("url", "")
            borrowed = source.get("borrowed", "")
            if url:
                lines.append(f"- {label}: {url}；借鉴点：{borrowed}")
            else:
                lines.append(f"- {label}；借鉴点：{borrowed}")
        lines.append("")
    lines.append("## 资料概况")
    lines.append("")
    lines.append(f"- 已读取文件：{len(package['documents'])}")
    lines.append(f"- 未解析/跳过文件：{len(package['unsupported_or_skipped'])}")
    lines.append(f"- 项目信号：{', '.join(package['signals']) or '未识别明显行业信号'}")
    lines.append("")
    lines.append("| 文档类型 | 数量 |")
    lines.append("| --- | ---: |")
    for key, value in sorted(package["document_type_counts"].items()):
        lines.append(f"| {escape_pipe(key)} | {value} |")
    if not package["document_type_counts"]:
        lines.append("| 未识别 | 0 |")
    lines.append("")
    lines.append("## 优先问题总表")
    lines.append("")
    lines.append("| ID | 优先级 | 置信度 | 主题 | 节点/对象 | 首要专家问题 |")
    lines.append("| --- | --- | --- | --- | --- | --- |")
    for issue in issues:
        first_q = issue["expert_questions"][0] if issue["expert_questions"] else ""
        lines.append(
            f"| {issue['issue_id']} | {PRIORITY_LABEL.get(issue['priority'], issue['priority'])} | {CONFIDENCE_LABEL.get(issue['confidence'], issue['confidence'])} | "
            f"{escape_pipe(issue['topic'])} | {escape_pipe(issue['node'])} | {escape_pipe(first_q)} |"
        )
    lines.append("")
    lines.append("## 详细专家意见")
    lines.append("")
    for issue in issues:
        lines.append(f"### {issue['issue_id']} {issue['topic']}")
        lines.append("")
        lines.append(f"- 优先级：{PRIORITY_LABEL.get(issue['priority'], issue['priority'])}")
        lines.append(f"- 置信度：{CONFIDENCE_LABEL.get(issue['confidence'], issue['confidence'])}（基于输入资料命中程度，不等同风险等级）")
        lines.append(f"- 节点/对象：{issue['node']}")
        lines.append(f"- 导向词：{', '.join(issue['guidewords']) or '未指定'}")
        lines.append(f"- 触发规则：`{issue['source_rule']}`")
        lines.append("")
        if issue["concern"]:
            lines.append("关注点：")
            lines.append("")
            lines.append(issue["concern"])
            lines.append("")
        if issue["expert_questions"]:
            lines.append("建议在会上提出：")
            lines.append("")
            for i, q in enumerate(issue["expert_questions"], start=1):
                lines.append(f"{i}. {q}")
            lines.append("")
        if issue["requested_evidence"]:
            lines.append("建议要求补充/定位的证据：")
            lines.append("")
            for item in issue["requested_evidence"]:
                lines.append(f"- {item}")
            lines.append("")
        if issue["evidence"]:
            lines.append("输入资料证据摘录：")
            lines.append("")
            for ev in issue["evidence"]:
                tags = f"；tags: {', '.join(ev['tags'])}" if ev.get("tags") else ""
                lines.append(f"- `{ev['rel_path']}`；关键词：`{ev['keyword']}`{tags}")
                lines.append(f"  - {ev['excerpt']}")
            lines.append("")
        else:
            lines.append("输入资料证据摘录：无直接摘录；该项主要由资料缺口或项目级规则触发。")
            lines.append("")
    lines.append("## 文件索引摘录")
    lines.append("")
    lines.append("| 文件 | 方法 | 类型 | 字符数 | 警告 |")
    lines.append("| --- | --- | --- | ---: | --- |")
    for doc in package["documents"][:80]:
        lines.append(
            f"| `{escape_pipe(doc['rel_path'])}` | {escape_pipe(doc['method'])} | "
            f"{escape_pipe(', '.join(doc['doc_types']))} | {doc['text_chars']} | {escape_pipe(', '.join(doc['warnings']))} |"
        )
    if len(package["documents"]) > 80:
        lines.append(f"| ... | ... | ... | ... | 另有 {len(package['documents']) - 80} 个文件见 JSON 索引 |")
    return "\n".join(lines) + "\n"


def render_drawing_reviews_markdown(package: dict[str, Any]) -> str:
    meta = package["metadata"]
    context = package.get("project_context", {})
    reviews = package.get("drawing_reviews", [])
    finding_count = sum(len(review["findings"]) for review in reviews)
    lines: list[str] = []
    lines.append(f"# HAZOP逐图纸专家预审表 - {meta['project_name']}")
    lines.append("")
    lines.append(f"- 生成时间：{meta['generated_at']}")
    lines.append(f"- 输入目录：`{meta['input_dir']}`")
    lines.append("- 使用角色：HAZOP参会专家会前预审，不是主席结论、正式worksheet或行动项关闭意见")
    lines.append(f"- 参与逐张分析的P&ID数：{len(reviews)}")
    lines.append(f"- 逐图纸问题行：{finding_count}")
    lines.append("")
    lines.append("## 使用口径")
    lines.append("")
    lines.append("本表按人类HAZOP审查习惯组织：逐张P&ID先确认节点/设计意图，再围绕参数和导向词提出问题、可能原因、可能后果、已有措施待核查和专家行动。")
    lines.append("跨图纸关系只在确有接口、联锁、公用工程、BOG、ESD/SIS等上下游关系时列为“必要时跨图核查”，不把全项目问题混成一张总表。")
    lines.append("所有“原因、后果、措施”均为会前待确认项，必须由设计方、运行、仪表、电气、消防、总图和厂家用正式资料确认。")
    lines.append("PFD只作为项目整体理解基础，不参与逐张问题行分析；文件名或路径含“已建”的图纸默认忽略。")
    if context:
        lines.append("")
        lines.append("## 项目整体理解基础")
        lines.append("")
        lines.append(f"- PFD背景图纸：{len(context.get('pfd_context_documents', []))} 份")
        for item in context.get("pfd_context_documents", [])[:8]:
            lines.append(f"- `{item.get('rel_path', '')}`：用于理解整体流程，不参与逐张P&ID分析。")
        lines.append(f"- 已建图纸默认忽略：{len(context.get('ignored_existing_drawings', []))} 份")
        lines.append(f"- 通用注释/图例/详图默认忽略：{len(context.get('ignored_reference_drawings', []))} 份")
    lines.append("")
    lines.append("## 图纸总览")
    lines.append("")
    lines.append("| 图纸ID | 图号 | P&ID图纸 | 节点提示 | 问题行 | 抽取字符 |")
    lines.append("| --- | --- | --- | --- | ---: | ---: |")
    for review in reviews:
        lines.append(
            f"| {review['drawing_id']} | {escape_pipe(review['drawing_no'])} | "
            f"{escape_pipe(review['title'])} | {escape_pipe(review['node_hint'])} | "
            f"{len(review['findings'])} | {review['text_chars']} |"
        )
    lines.append("")
    lines.append("## 逐图纸预审")
    lines.append("")
    for review in reviews:
        lines.append(f"### {review['drawing_id']} {review['drawing_no']} {review['title']}")
        lines.append("")
        lines.append(f"- 文件：`{review['rel_path']}`")
        lines.append(f"- 节点提示：{review['node_hint']}")
        lines.append(f"- 设计意图提示：{review['design_intent_hint']}")
        lines.append(f"- 审查关注：{'；'.join(review['review_focus'])}")
        if review["extraction_warnings"]:
            lines.append(f"- 抽取警告：{'；'.join(review['extraction_warnings'])}")
        lines.append("")
        if not review["findings"]:
            lines.append("本图纸未形成自动问题行。专家仍需人工看图确认节点边界、阀位、跨图连接和仪表联锁。")
            lines.append("")
            continue
        for finding in review["findings"]:
            lines.append(f"#### {finding['finding_id']} {finding['parameter']} / {PRIORITY_LABEL.get(finding['priority'], finding['priority'])}")
            lines.append("")
            lines.append(f"- 分析对象：{finding['analysis_object']}")
            lines.append(f"- 导向词：{', '.join(finding['guidewords']) or '未指定'}")
            lines.append(f"- 置信度：{CONFIDENCE_LABEL.get(finding['confidence'], finding['confidence'])}")
            lines.append(f"- 触发规则：`{finding['source_rule']}`")
            lines.append("")
            lines.append("问题/偏差：")
            lines.append("")
            lines.append(f"- {finding['problem']}")
            lines.append("")
            lines.append("可能原因（待确认）：")
            lines.append("")
            for item in finding["possible_causes"]:
                lines.append(f"- {item}")
            lines.append("")
            lines.append("可能后果（待确认）：")
            lines.append("")
            for item in finding["possible_consequences"]:
                lines.append(f"- {item}")
            lines.append("")
            lines.append("已有措施/保护层待核查：")
            lines.append("")
            for item in finding["existing_safeguards_to_verify"]:
                lines.append(f"- {item}")
            lines.append("")
            lines.append("专家行动/需补证据：")
            lines.append("")
            for item in finding["expert_actions"]:
                lines.append(f"- {item}")
            if finding["related_documents"]:
                lines.append("")
                lines.append("必要时跨图核查：")
                lines.append("")
                for item in finding["related_documents"]:
                    lines.append(f"- `{item}`")
            if finding["evidence"]:
                lines.append("")
                lines.append("本图纸证据摘录：")
                lines.append("")
                for ev in finding["evidence"]:
                    tags = f"；tags: {', '.join(ev['tags'])}" if ev.get("tags") else ""
                    lines.append(f"- 关键词：`{ev['keyword']}`{tags}")
                    lines.append(f"  - {ev['excerpt']}")
            lines.append("")
    return "\n".join(lines) + "\n"


def write_csv(path: Path, issues: list[dict[str, Any]]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "issue_id", "priority", "confidence", "topic", "node", "guidewords",
                "concern", "expert_questions", "requested_evidence", "evidence_files", "source_rule"
            ],
        )
        writer.writeheader()
        for issue in issues:
            writer.writerow({
                "issue_id": issue["issue_id"],
                "priority": PRIORITY_LABEL.get(issue["priority"], issue["priority"]),
                "confidence": CONFIDENCE_LABEL.get(issue["confidence"], issue["confidence"]),
                "topic": issue["topic"],
                "node": issue["node"],
                "guidewords": "; ".join(issue["guidewords"]),
                "concern": issue["concern"],
                "expert_questions": "\n".join(issue["expert_questions"]),
                "requested_evidence": "; ".join(issue["requested_evidence"]),
                "evidence_files": "; ".join(sorted({ev["rel_path"] for ev in issue["evidence"]})),
                "source_rule": issue["source_rule"],
            })


def write_drawing_review_csv(path: Path, reviews: list[dict[str, Any]]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "drawing_id", "drawing_no", "title", "rel_path", "node_hint", "design_intent_hint",
                "finding_id", "priority", "confidence", "parameter", "guidewords", "analysis_object",
                "problem", "possible_causes", "possible_consequences", "existing_safeguards_to_verify",
                "expert_actions", "related_documents", "evidence", "source_rule",
            ],
        )
        writer.writeheader()
        for review in reviews:
            if not review["findings"]:
                writer.writerow({
                    "drawing_id": review["drawing_id"],
                    "drawing_no": review["drawing_no"],
                    "title": review["title"],
                    "rel_path": review["rel_path"],
                    "node_hint": review["node_hint"],
                    "design_intent_hint": review["design_intent_hint"],
                    "finding_id": "",
                    "priority": "",
                    "confidence": "",
                    "parameter": "",
                    "guidewords": "",
                    "analysis_object": "",
                    "problem": "未形成自动问题行，需人工看图确认。",
                    "possible_causes": "",
                    "possible_consequences": "",
                    "existing_safeguards_to_verify": "",
                    "expert_actions": "人工确认节点边界、阀位、跨图连接和仪表联锁。",
                    "related_documents": "",
                    "evidence": "",
                    "source_rule": "",
                })
                continue
            for finding in review["findings"]:
                writer.writerow({
                    "drawing_id": review["drawing_id"],
                    "drawing_no": review["drawing_no"],
                    "title": review["title"],
                    "rel_path": review["rel_path"],
                    "node_hint": review["node_hint"],
                    "design_intent_hint": review["design_intent_hint"],
                    "finding_id": finding["finding_id"],
                    "priority": PRIORITY_LABEL.get(finding["priority"], finding["priority"]),
                    "confidence": CONFIDENCE_LABEL.get(finding["confidence"], finding["confidence"]),
                    "parameter": finding["parameter"],
                    "guidewords": "; ".join(finding["guidewords"]),
                    "analysis_object": finding["analysis_object"],
                    "problem": finding["problem"],
                    "possible_causes": "\n".join(finding["possible_causes"]),
                    "possible_consequences": "\n".join(finding["possible_consequences"]),
                    "existing_safeguards_to_verify": "\n".join(finding["existing_safeguards_to_verify"]),
                    "expert_actions": "\n".join(finding["expert_actions"]),
                    "related_documents": "\n".join(finding["related_documents"]),
                    "evidence": "\n".join(f"{ev['keyword']}: {ev['excerpt']}" for ev in finding["evidence"]),
                    "source_rule": finding["source_rule"],
                })


def render_review_prompt(package: dict[str, Any]) -> str:
    meta = package["metadata"]
    top = package["issues"][:20]
    drawing_reviews = package.get("drawing_reviews", [])
    drawing_focus = []
    for review in drawing_reviews:
        if review["findings"]:
            drawing_focus.append(review)
        if len(drawing_focus) >= 12:
            break
    lines = [
        f"# 复核提示包 - {meta['project_name']}",
        "",
        "你是HAZOP参会专家，不是主席。请优先基于 `hazop_drawing_reviews.md/json` 和源设计文件逐张图纸复核；`hazop_expert_opinions.md/json` 只作为规则命中和资料缺口的背景池。",
        "本轮逐张复核只针对新建/扩建P&ID。PFD仅用于理解项目整体流程和边界；文件名或路径含“已建”的图纸默认作为既有装置接口背景，不进入逐张问题行复核。",
        "",
        "1. 一次只复核一张图纸，先确认节点边界、设计意图和分析对象。",
        "2. 对每个图内问题行核查：问题/偏差、可能原因、可能后果、已有措施、需补证据是否与图纸一致。",
        "3. 删除没有证据或不适用于本图纸边界的问题行。",
        "4. 合并同一图纸内重复问题，保留最具体、最可执行的专家发问。",
        "5. 只有涉及接口、上下游、ESD/SIS/C&E、BOG、火炬/放空、公用工程或共享保护层时，才跨图纸描述前后关系。",
        "6. 不要给出主席结论，不要关闭行动项。",
        "",
        "## 优先复核图纸",
        "",
    ]
    for review in drawing_focus:
        first = review["findings"][0]
        lines.append(
            f"- {review['drawing_id']} {review['title']}（{review['node_hint']}）："
            f"{first['finding_id']} [{PRIORITY_LABEL.get(first['priority'], first['priority'])}] {first['problem']}"
        )
    lines.extend([
        "",
        "## 全局规则池抽查",
        "",
    ])
    for issue in top:
        q = issue["expert_questions"][0] if issue["expert_questions"] else ""
        lines.append(f"- {issue['issue_id']} [{PRIORITY_LABEL.get(issue['priority'], issue['priority'])}] {issue['topic']}: {q}")
    return "\n".join(lines) + "\n"


def escape_pipe(text: Any) -> str:
    return str(text).replace("|", "\\|").replace("\n", " ").strip()


if __name__ == "__main__":
    raise SystemExit(main())
